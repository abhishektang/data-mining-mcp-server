#!/usr/bin/env python3
"""
Fast text extraction from PDF and DOCX files using PyMuPDF and python-docx
No OCR - extracts embedded text directly (much faster!)
"""
import sys
import json
from pathlib import Path

def extract_pdf_text(file_path: str, max_pages: int = None) -> dict:
    """Extract text from PDF using PyMuPDF (very fast!)"""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages_to_extract = min(total_pages, max_pages) if max_pages else total_pages
        
        text_parts = []
        for page_num in range(pages_to_extract):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(f"## Page {page_num + 1}\n\n{text.strip()}")
        
        doc.close()
        
        content = "\n\n---\n\n".join(text_parts) if text_parts else "[No text found in PDF]"
        
        return {
            "success": True,
            "content": content,
            "total_pages": total_pages,
            "pages_extracted": pages_to_extract,
            "file_type": "PDF"
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"PDF extraction error: {str(e)}"
        }


def extract_docx_text(file_path: str) -> dict:
    """Extract text from DOCX using python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        content = "\n\n".join(text_parts) if text_parts else "[No text found in DOCX]"
        
        return {
            "success": True,
            "content": content,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "file_type": "DOCX"
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"DOCX extraction error: {str(e)}"
        }


def extract_text(file_path: str, max_pages: int = None) -> dict:
    """
    Extract text from PDF or DOCX file.
    
    Args:
        file_path: Path to the file
        max_pages: Maximum pages to extract from PDF (None = all pages)
    
    Returns:
        dict with success, content, and metadata
    """
    path = Path(file_path)
    
    if not path.exists():
        return {
            "success": False,
            "error": f"File not found: {file_path}"
        }
    
    extension = path.suffix.lower()
    
    if extension == '.pdf':
        return extract_pdf_text(file_path, max_pages)
    elif extension in ['.docx', '.doc']:
        return extract_docx_text(file_path)
    else:
        return {
            "success": False,
            "error": f"Unsupported file type: {extension}"
        }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({
            "success": False,
            "error": "Usage: extract_text.py <file_path> [max_pages]"
        }))
        sys.exit(1)
    
    file_path = sys.argv[1]
    max_pages = int(sys.argv[2]) if len(sys.argv) > 2 else None
    
    result = extract_text(file_path, max_pages)
    print(json.dumps(result, indent=2))
    sys.exit(0 if result["success"] else 1)
